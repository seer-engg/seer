"""Document processor service for extracting text from various file formats."""
from __future__ import annotations

import io
from typing import Dict, Set

from seer.logger import get_logger

logger = get_logger("services.knowledge.document_processor")


# Supported MIME types
SUPPORTED_MIME_TYPES: Set[str] = {
    "text/plain",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
}

# MIME type to file extension mapping
MIME_TO_EXTENSION: Dict[str, str] = {
    "text/plain": ".txt",
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
}


class DocumentProcessor:
    """Service for extracting text content from uploaded documents."""

    @staticmethod
    def is_supported(mime_type: str) -> bool:
        """Check if a MIME type is supported.

        Args:
            mime_type: MIME type to check

        Returns:
            True if the MIME type is supported
        """
        return mime_type in SUPPORTED_MIME_TYPES

    async def extract_text(self, content: bytes, mime_type: str) -> str:
        """Extract text content from a document.

        Args:
            content: Raw file content as bytes
            mime_type: MIME type of the file

        Returns:
            Extracted text content

        Raises:
            ValueError: If MIME type is not supported
            Exception: If text extraction fails
        """
        if not self.is_supported(mime_type):
            raise ValueError(f"Unsupported MIME type: {mime_type}. Supported: {', '.join(SUPPORTED_MIME_TYPES)}")

        if mime_type == "text/plain":
            return self._extract_plain_text(content)
        if mime_type == "application/pdf":
            return await self._extract_pdf_text(content)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await self._extract_docx_text(content)
        raise ValueError(f"Unsupported MIME type: {mime_type}")

    def _extract_plain_text(self, content: bytes) -> str:
        """Extract text from plain text file.

        Args:
            content: Raw file content

        Returns:
            Decoded text content
        """
        # Try UTF-8 first, then fallback to latin-1
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            logger.warning("UTF-8 decode failed, trying latin-1")
            return content.decode("latin-1")

    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF file using pypdf.

        Args:
            content: Raw PDF content

        Returns:
            Extracted text content
        """
        try:
            from pypdf import PdfReader  # pylint: disable=import-outside-toplevel  # Reason: Lazy import for optional dependency
        except ImportError as exc:
            raise ImportError("pypdf is required for PDF processing. Install with: pip install pypdf") from exc

        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:  # pylint: disable=broad-exception-caught  # Reason: Catch all PDF parsing errors
                logger.warning("Failed to extract text from page %d: %s", page_num, e)

        text = "\n\n".join(text_parts)
        logger.debug("Extracted PDF text", extra={"pages": len(reader.pages), "text_length": len(text)})
        return text

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file using python-docx.

        Args:
            content: Raw DOCX content

        Returns:
            Extracted text content
        """
        try:
            from docx import Document  # pylint: disable=import-outside-toplevel  # Reason: Lazy import for optional dependency
        except ImportError as exc:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx") from exc

        docx_file = io.BytesIO(content)
        doc = Document(docx_file)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        text = "\n\n".join(text_parts)
        logger.debug("Extracted DOCX text", extra={"paragraphs": len(doc.paragraphs), "text_length": len(text)})
        return text


# Singleton instance
_DOCUMENT_PROCESSOR: DocumentProcessor | None = None


def get_document_processor() -> DocumentProcessor:
    """Get or create singleton document processor."""
    global _DOCUMENT_PROCESSOR  # pylint: disable=global-statement  # Reason: Singleton pattern for service instance
    if _DOCUMENT_PROCESSOR is None:
        _DOCUMENT_PROCESSOR = DocumentProcessor()
    return _DOCUMENT_PROCESSOR


__all__ = [
    "DocumentProcessor",
    "get_document_processor",
    "SUPPORTED_MIME_TYPES",
    "MIME_TO_EXTENSION",
]
